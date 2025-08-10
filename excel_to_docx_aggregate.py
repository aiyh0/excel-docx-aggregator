#!/usr/bin/env python3
"""
excel_to_docx_aggregate.py  (strict header matching + optional Unit Code)

- Stricter synonyms (reduced variants)
- Higher fuzzy MATCH_THRESHOLD to reduce false positives
- Optional "Unit Code" column (e.g., UOM) included before Description if present

Usage:
  python excel_to_docx_aggregate.py input.xlsx -o output.docx
  # Auto-detect newest Excel in a folder:
  python excel_to_docx_aggregate.py --dir .
"""

import argparse
import math
import re
import sys
import os
import glob
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import numpy as np
import pandas as pd

# Prefer rapidfuzz; fall back to fuzzywuzzy if needed
try:
    from rapidfuzz import fuzz, process
    RF_AVAILABLE = True
except Exception:
    RF_AVAILABLE = False
    try:
        from fuzzywuzzy import fuzz, process  # type: ignore
    except Exception as e:
        print("Error: Install 'rapidfuzz' (preferred) or 'fuzzywuzzy' for fuzzy matching.")
        raise e

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------- Configurable bits ----------
SCAN_ROWS = 10  # how many rows (from top) to scan for header candidates
MATCH_THRESHOLD = 88  # stricter fuzzy threshold (was 70)

# Strict(er) synonyms / small-typo variants only
COL_SYNONYMS = {
    "description": {"description", "desc"},
    "quantity": {"quantity", "qty"},
    "unit price": {"unit price", "unit cost", "unitprice", "unit-cost"},
    "line amount": {"line amount", "line total", "extended price", "ext price"},
    "unit code": {"unit code", "uom", "u/m", "uom code"},
}

DOCX_TABLE_STYLE_PREFS = ["Light List Accent 1", "Light Shading Accent 1", "Table Grid"]


# ---------- Helpers ----------
def norm_text(s: str) -> str:
    if s is None:
        return ""
    s = str(s)
    s = s.strip().lower()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[_\-:/\\]+", " ", s)  # normalize common separators
    return s


def best_match(label: str, choices: list[str]) -> tuple[str, int]:
    """Return (choice, score). Uses WRatio/ratio-like depending on backend."""
    label_n = norm_text(label)
    choices_n = [norm_text(c) for c in choices]
    if RF_AVAILABLE:
        res = process.extractOne(label_n, choices_n, scorer=fuzz.WRatio, score_cutoff=0)
        if res:
            return choices[res[2]], int(res[1])
        return "", 0
    else:
        best_c, best_s = "", 0
        from typing import Iterable  # noqa
        for ch_raw, ch in zip(choices, choices_n):
            s = fuzz.WRatio(label_n, ch)
            if s > best_s:
                best_c, best_s = ch_raw, s
        return best_c, best_s


def row_header_map(cells: list[str]) -> tuple[dict[int, str], int]:
    """
    Try to map each required semantic column to a column index in this row.
    Returns (mapping: col_index -> semantic_name, score_sum)
    """
    mapping: dict[int, str] = {}
    score_sum = 0

    semantic_targets = []
    for semantic, syns in COL_SYNONYMS.items():
        semantic_targets.append((semantic, list(syns)))

    used_cols = set()
    for semantic, choices in semantic_targets:
        best_idx, best_score = -1, -1
        for idx, cell in enumerate(cells):
            if idx in used_cols:
                continue
            _choice, score = best_match(cell, choices)
            if score > best_score:
                best_score = score
                best_idx = idx
        if best_idx >= 0 and best_score >= MATCH_THRESHOLD:
            mapping[best_idx] = semantic
            score_sum += best_score
            used_cols.add(best_idx)

    return mapping, score_sum


def detect_header_row(df: pd.DataFrame) -> tuple[int | None, dict[int, str], int]:
    """
    Scan the top SCAN_ROWS for a header row. Returns (row_index, mapping, score)
    mapping is {col_index: 'description'|'quantity'|'unit price'|'line amount'|'unit code'}
    """
    best_row, best_map, best_score = None, {}, -1
    limit = min(SCAN_ROWS, len(df))

    for r in range(limit):
        cells = [(str(x) if not pd.isna(x) else "") for x in df.iloc[r, :].tolist()]
        mapping, score = row_header_map(cells)

        sems = set(mapping.values())
        has_desc = "description" in sems
        has_qty = "quantity" in sems
        has_line = "line amount" in sems

        if has_desc and (has_line or has_qty):
            if score > best_score:
                best_row, best_map, best_score = r, mapping, score

    return best_row, best_map, best_score


def coerce_numeric(series: pd.Series) -> pd.Series:
    s = pd.to_numeric(series, errors="coerce")
    return s


def safe_divide(numer: pd.Series, denom: pd.Series) -> pd.Series:
    """Divide numer by denom; convert +/-inf to NaN to avoid deprecated options."""
    with np.errstate(divide='ignore', invalid='ignore'):
        res = numer / denom
    res = res.replace([np.inf, -np.inf], np.nan)
    return res


def auto_find_excel(search_dir: str = ".") -> str:
    """Return the newest .xlsx/.xlsm file in search_dir (ignores temp files like '~$...')."""
    import glob, os
    patterns = ["*.xlsx", "*.xlsm"]
    files: list[str] = []
    for pat in patterns:
        files.extend(glob.glob(os.path.join(search_dir, pat)))
    files = [f for f in files if not os.path.basename(f).startswith("~$")]
    if not files:
        raise FileNotFoundError(f"No Excel files (*.xlsx, *.xlsm) found in {os.path.abspath(search_dir)}")
    files.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return files[0]


def pick_best_sheet(xl: pd.ExcelFile) -> tuple[str, int | None, dict[int, str], int]:
    best = ("", None, {}, -1)  # (sheet, header_row, mapping, score)
    for sheet in xl.sheet_names:
        df = xl.parse(sheet_name=sheet, header=None, dtype=object)
        row, mapping, score = detect_header_row(df)
        if score > best[3]:
            best = (sheet, row, mapping, score)
    return best


def standardize_columns(df: pd.DataFrame, mapping: dict[int, str]) -> pd.DataFrame:
    new_cols = []
    for i, c in enumerate(df.columns):
        new_cols.append(mapping.get(i, str(c)))
    df.columns = new_cols
    return df


def compute_missing_amounts(df: pd.DataFrame) -> pd.DataFrame:
    qty_col = "quantity" if "quantity" in df.columns else None
    price_col = "unit price" if "unit price" in df.columns else None
    amt_col = "line amount" if "line amount" in df.columns else None

    if qty_col:
        df[qty_col] = coerce_numeric(df[qty_col])
    if price_col:
        df[price_col] = coerce_numeric(df[price_col])
    if amt_col:
        df[amt_col] = coerce_numeric(df[amt_col])

    # Compute line amount if missing
    if amt_col is None and qty_col and price_col:
        df["line amount"] = df[qty_col] * df[price_col]
        amt_col = "line amount"

    # Compute unit price if missing (from amount and qty)
    if price_col is None and amt_col and qty_col:
        df["unit price"] = safe_divide(df[amt_col], df[qty_col])
        price_col = "unit price"

    return df


def aggregate_by_description(df: pd.DataFrame) -> pd.DataFrame:
    if "description" not in df.columns:
        raise ValueError("Could not find a Description column after header detection.")

    # Clean description
    df["__desc_raw__"] = df["description"].astype(str).fillna("").map(lambda x: x.strip())
    df["__desc_key__"] = df["__desc_raw__"].str.lower()

    # Optional Unit Code
    has_unit = "unit code" in df.columns
    if has_unit:
        df["__unit_raw__"] = df["unit code"].astype(str).fillna("").map(lambda x: x.strip())
        df["__unit_key__"] = df["__unit_raw__"].str.lower()

    # Numeric coercions and computations
    df = compute_missing_amounts(df)

    # Drop empty description rows
    df = df[df["__desc_raw__"].str.len() > 0].copy()

    # Build aggregation dict
    agg_dict = {}
    if "quantity" in df.columns:
        agg_dict["quantity"] = "sum"
    if "line amount" in df.columns:
        agg_dict["line amount"] = "sum"
    if not agg_dict:
        raise ValueError("No numeric columns (Quantity or Line Amount) found to aggregate.")

    # Group keys
    group_keys = ["__desc_key__"]
    if has_unit:
        group_keys = ["__unit_key__", "__desc_key__"]

    grouped = df.groupby(group_keys, dropna=True).agg(agg_dict).reset_index()

    # Attach pretty display values (first seen)
    pick_cols = group_keys + ["__desc_raw__"] + (["__unit_raw__"] if has_unit else [])
    first_seen = df.drop_duplicates(group_keys)[pick_cols]
    grouped = grouped.merge(first_seen, on=group_keys, how="left")

    # Weighted unit price if both Quantity & Line Amount present
    if ("quantity" in grouped.columns) and ("line amount" in grouped.columns):
        grouped["unit price (weighted)"] = safe_divide(grouped["line amount"], grouped["quantity"])

    # Order columns: Unit Code (if any), Description, then numeric
    out_cols = []
    if has_unit:
        out_cols.append("__unit_raw__")
    out_cols += ["__desc_raw__"]
    if "quantity" in grouped.columns:
        out_cols.append("quantity")
    if "unit price (weighted)" in grouped.columns:
        out_cols.append("unit price (weighted)")
    if "line amount" in grouped.columns:
        out_cols.append("line amount")

    grouped = grouped[out_cols].rename(columns={
        "__unit_raw__": "Unit Code",
        "__desc_raw__": "Description",
        "quantity": "Quantity",
        "line amount": "Line Amount",
        "unit price (weighted)": "Unit Price"
    })

    # Sort by Quantity asc, then Unit Code/Description
    sort_cols = []
    if "Quantity" in grouped.columns:
        sort_cols.append("Quantity")
    if "Unit Code" in grouped.columns:
        sort_cols.append("Unit Code")
    sort_cols.append("Description")
    grouped = grouped.sort_values(sort_cols, ascending=[True]*len(sort_cols), na_position="last")

    # Ensure numeric dtype
    for col in ["Quantity", "Unit Price", "Line Amount"]:
        if col in grouped.columns:
            grouped[col] = pd.to_numeric(grouped[col], errors="coerce")

    # Append TOTAL row (Unit Code blank)
    total = {"Description": "TOTAL"}
    if "Unit Code" in grouped.columns:
        total["Unit Code"] = ""
    if "Quantity" in grouped.columns:
        total["Quantity"] = grouped["Quantity"].sum(min_count=1)
    if "Line Amount" in grouped.columns:
        total["Line Amount"] = grouped["Line Amount"].sum(min_count=1)
    grouped = pd.concat([grouped, pd.DataFrame([total])], ignore_index=True)

    return grouped.reset_index(drop=True)


def export_to_docx(df: pd.DataFrame, out_path: str, title: Optional[str] = None):
    doc = Document()

    if title:
        h = doc.add_heading(title, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)

    # Style
    for style_name in ["Light List Accent 1", "Light Shading Accent 1", "Table Grid"]:
        try:
            table.style = style_name
            break
        except Exception:
            continue

    # Header row
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(str(col))
        run.bold = True

    # Data rows
    def fmt(x):
        if isinstance(x, (int, float)) and not (isinstance(x, float) and (math.isnan(x) or math.isinf(x))):
            return f"{x:,.2f}" if abs(x) >= 1e-6 else "0.00"
        return "" if x is None or (isinstance(x, float) and math.isnan(x)) else str(x)

    for i in range(rows):
        for j in range(cols):
            table.rows[i + 1].cells[j].text = fmt(df.iat[i, j])

    # Optional: set column widths (rough)
    try:
        if cols >= 1:
            table.columns[0].width = Inches(1.2 if df.columns[0] == "Unit Code" else 3.5)
        for j in range(1, cols):
            table.columns[j].width = Inches(1.3)
    except Exception:
        pass

    doc.save(out_path)


def load_and_process(input_path: str, sheet: Optional[str] = None) -> pd.DataFrame:
    xl = pd.ExcelFile(input_path, engine="openpyxl")
    if sheet is None:
        best_sheet, header_row, mapping, score = pick_best_sheet(xl)
        if header_row is None or score < 0:
            raise ValueError("Could not detect headers in the first few rows of any sheet.")
        print(f"[Info] Auto-selected sheet: {best_sheet} (header row ~ {header_row+1})")
        df_raw = xl.parse(sheet_name=best_sheet, header=None, dtype=object)
    else:
        df_raw = xl.parse(sheet_name=sheet, header=None, dtype=object)
        header_row, mapping, score = detect_header_row(df_raw)
        if header_row is None:
            raise ValueError(f"Could not detect headers in the first {SCAN_ROWS} rows of sheet '{sheet}'.")
        print(f"[Info] Using sheet: {sheet} (header row ~ {header_row+1})")

    # Re-read with proper header row so pandas assigns names, then remap to semantic names
    df = df_raw.copy()
    df.columns = df.iloc[header_row]
    df = df.iloc[header_row + 1:].reset_index(drop=True)

    # Build mapping again on the header row text
    header_cells = [str(x) if not pd.isna(x) else "" for x in df_raw.iloc[header_row, :].tolist()]
    mapping, _ = row_header_map(header_cells)

    # Keep only columns from header row width
    df = df.iloc[:, :len(header_cells)].copy()
    df = standardize_columns(df, mapping)

    # Only keep relevant columns + anything we need for computations
    keep_cols = {"description", "quantity", "unit price", "line amount", "unit code"}
    existing_relevant = [c for c in df.columns if c in keep_cols]
    if "description" not in existing_relevant:
        raise ValueError("Could not find a Description-like column after detection.")
    df = df[existing_relevant].copy()

    # Clean up description a bit
    df["description"] = df["description"].apply(lambda v: "" if pd.isna(v) else " ".join(str(v).strip().split()))

    return df


def main():
    parser = argparse.ArgumentParser(description="Extract, group, and export Excel line items to a DOCX table.")
    parser.add_argument("input", nargs="?", help="Path to the input .xlsx/.xlsm file (optional; auto-detects newest if omitted)")
    parser.add_argument("-o", "--output", default=None, help="Path to the output .docx file (defaults to <input>_aggregated.docx)")
    parser.add_argument("--dir", default=None, help="Directory to auto-detect Excel file from (defaults to current working directory)")
    parser.add_argument("--sheet", default=None, help="Optional sheet name to use")
    parser.add_argument("--title", default="Aggregated Line Items", help="Heading title in the docx")
    args = parser.parse_args()

    try:
        # Determine input path (auto-detect newest Excel if not provided)
        input_path = args.input
        if input_path is None:
            search_dir = args.dir if args.dir else os.getcwd()
            input_path = auto_find_excel(search_dir)
            print(f"[Info] Auto-selected file: {input_path}")

        # Determine output path
        if args.output:
            output_path = args.output
        else:
            p = Path(input_path)
            output_path = str(p.with_name(f"{p.stem}_aggregated.docx"))

        df = load_and_process(input_path, sheet=args.sheet)
        agg = aggregate_by_description(df)
        if agg.empty:
            print("No rows to export after aggregation.")
            sys.exit(2)
        export_to_docx(agg, output_path, title=args.title)
        print(f"Done. Wrote {len(agg)} rows to: {output_path}")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
